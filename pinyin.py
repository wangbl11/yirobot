# coding=utf-8
from xpinyin import Pinyin
import sys
import re
import math
from docx import Document

filename = sys.argv[1]
outfile = sys.argv[2]
sep = sys.argv[3]
pat = re.compile(r'[\s\n\r]')
p = Pinyin()
lensize = 0
with open(filename) as file_obj:
    lines = file_obj.readlines()
    lensize = len(lines)
titles = ['一', '二', '三', '四', '五', '六', '七', '八', '九']


def one_chapter(lines, start, end, idx):
    outs = []
    document.add_heading('单元' + titles[idx], level=2)
    # with open(outfile,'w') as file_out:
    for i in range(start, end):
        line = lines[i]
        line = pat.sub('', line)
        outs.extend(line.split(sep))
        # file_out.write(p.get_pinyin(line, show_tone_marks=True))

    outs = list(filter(lambda n: ('##' not in n) and n != '', outs))

    _len = len(outs)
    print(_len)
    cnt = math.ceil(_len / 4)
    table = document.add_table(rows=cnt, cols=4, style='Table Grid')

    idx = 0
    for i in range(0, cnt):
        hdr_cells = table.rows[i].cells
        for j in range(4):
            if idx < _len:
                hdr_cells[j].text = p.get_pinyin(outs[idx], tone_marks='marks') + '\n\n'
                idx = idx + 1
            else:
                break
    document.add_page_break()


document = Document()

start = 0
indexes = [i for i, x in enumerate(lines) if '##' in x]
cnt = len(indexes)
print(indexes)

for j in range(0, cnt):
    start = indexes[j]
    # if indexes[j] == 0:
    #     continue
    one_chapter(lines, start, lensize if j == cnt - 1 else indexes[j + 1], j)


document.save(outfile)
